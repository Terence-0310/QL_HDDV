# -*- coding: utf-8 -*-
"""
Convert BAO_CAO_DO_AN_ECMS.md -> a formatted academic DOCX report
matching the DACS (Do An Co So) report layout (like NHOM2_Final3.docx).
"""
import re
import sys
import os
import struct
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"d:\Hutech\Đồ Án\Hệ thống quản lý hợp đồng điện tử và nhắc mốc gia hạn"
SRC = BASE + r"\BAO_CAO_DO_AN_ECMS.md"
OUT = BASE + r"\BAO_CAO_DO_AN_ECMS.docx"
DIAG = BASE + r"\assets\diagrams"

MAX_W_CM = 15.5   # content width
MAX_H_CM = 19.5   # usable height

# Caption titles for the 8 Mermaid diagrams (in document order)
DIAGRAM_TITLES = [
    "Kiến trúc tổng quan hệ thống ECMS",
    "Kiến trúc phân tầng (Layered Architecture)",
    "Sơ đồ thực thể quan hệ (ERD)",
    "Sơ đồ Use Case tổng quát",
    "Sequence Diagram – Đăng nhập hệ thống",
    "Sequence Diagram – Tạo hợp đồng",
    "Sequence Diagram – Phê duyệt hợp đồng",
    "Sơ đồ quy trình chuẩn hóa dữ liệu",
]


def png_size(path):
    with open(path, 'rb') as f:
        head = f.read(24)
    if head[:8] != b'\x89PNG\r\n\x1a\n':
        return None
    w, h = struct.unpack('>II', head[16:24])
    return w, h

FONT = "Times New Roman"
BODY_SIZE = 13
HEAD_COLOR = RGBColor(0x00, 0x00, 0x00)   # all black & white
BLACK = RGBColor(0x00, 0x00, 0x00)
TBL_HEADER_FILL = 'D9D9D9'                 # light gray header (grayscale)

# ---------------------------------------------------------------- helpers
def set_run_font(run, size=BODY_SIZE, bold=False, italic=False, color=None, mono=False):
    name = "Consolas" if mono else FONT
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), name)


def shade_cell(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcpr.append(shd)


def shade_paragraph(p, hexcolor):
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    ppr.append(shd)


def add_box_border(p):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '6')
        e.set(qn('w:space'), '6')
        e.set(qn('w:color'), 'BBBBBB')
        pbdr.append(e)
    ppr.append(pbdr)


INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))')
TAG_RE = re.compile(r'</?[a-zA-Z][^>]*>')


def clean_inline_text(t):
    t = TAG_RE.sub('', t)
    t = t.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&nbsp;', ' ')
    return t


def add_inline(p, text, base_size=BODY_SIZE, base_bold=False, base_italic=False, base_color=None):
    """Parse markdown inline formatting and append runs to paragraph p."""
    text = clean_inline_text(text)
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            set_run_font(r, base_size, base_bold, base_italic, base_color)
        tok = m.group(0)
        if tok.startswith('**'):
            r = p.add_run(tok[2:-2])
            set_run_font(r, base_size, True, base_italic, base_color)
        elif tok.startswith('`'):
            r = p.add_run(tok[1:-1])
            set_run_font(r, base_size - 1, base_bold, base_italic, base_color, mono=True)
        elif tok.startswith('['):
            lt = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok)
            r = p.add_run(lt.group(1))
            set_run_font(r, base_size, base_bold, base_italic, base_color)
        elif tok.startswith('*'):
            r = p.add_run(tok[1:-1])
            set_run_font(r, base_size, base_bold, True, base_color)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        set_run_font(r, base_size, base_bold, base_italic, base_color)


# ---------------------------------------------------------------- document base
doc = Document()

# base style
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(BODY_SIZE)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.5
pf.space_after = Pt(6)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for sec in doc.sections:
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)

# heading styles
for lvl, sz in ((1, 16), (2, 14), (3, 13), (4, 13)):
    st = doc.styles['Heading %d' % lvl]
    st.font.name = FONT
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = HEAD_COLOR
    st.font.italic = (lvl == 4)
    st.paragraph_format.space_before = Pt(12 if lvl <= 2 else 6)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.3
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def heading(text, level, page_break=False):
    if page_break:
        doc.add_page_break()
    h = doc.add_heading(level=level)
    r = h.add_run(clean_inline_text(text))
    set_run_font(r, {1: 16, 2: 14, 3: 13, 4: 13}[level], True, level == 4, HEAD_COLOR)
    return h


def para(text='', align=None, size=BODY_SIZE, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_inline(p, text, size, bold, italic, color)
    return p


# ---------------------------------------------------------------- title page
def build_title_page():
    para('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HỒ CHÍ MINH (HUTECH)',
         WD_ALIGN_PARAGRAPH.CENTER, 13, True, space_after=0)
    para('KHOA CÔNG NGHỆ THÔNG TIN', WD_ALIGN_PARAGRAPH.CENTER, 13, True, space_after=0)
    para('-----o0o-----', WD_ALIGN_PARAGRAPH.CENTER, 13, space_after=24)

    para('BÁO CÁO ĐỒ ÁN CƠ SỞ', WD_ALIGN_PARAGRAPH.CENTER, 18, True,
         color=BLACK, space_after=18)

    para('ĐỀ TÀI:', WD_ALIGN_PARAGRAPH.CENTER, 13, True, space_after=6)
    para('HỆ THỐNG QUẢN LÝ HỢP ĐỒNG ĐIỆN TỬ\nVÀ NHẮC MỐC GIA HẠN',
         WD_ALIGN_PARAGRAPH.CENTER, 20, True, color=BLACK,
         space_after=6)
    para('(Electronic Contract Management & Automated Reminder System)',
         WD_ALIGN_PARAGRAPH.CENTER, 13, italic=True, space_after=36)

    for label, val in (
        ('Ngành:', 'CÔNG NGHỆ THÔNG TIN'),
        ('Chuyên ngành:', 'Kỹ thuật phần mềm'),
    ):
        p = para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        r = p.add_run(label + ' '); set_run_font(r, 13, False)
        r = p.add_run(val); set_run_font(r, 13, True)
    doc.add_paragraph()

    for label, val in (
        ('Giảng viên hướng dẫn:', '............................................'),
        ('Sinh viên thực hiện:', '............................................'),
        ('MSSV:', '............................................'),
        ('Lớp:', '............................................'),
    ):
        p = para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        r = p.add_run(label + ' '); set_run_font(r, 13, False)
        r = p.add_run(val); set_run_font(r, 13, True)

    doc.add_paragraph(); doc.add_paragraph()
    para('TP. Hồ Chí Minh, tháng 06 năm 2026',
         WD_ALIGN_PARAGRAPH.CENTER, 13, True, italic=True)


# ---------------------------------------------------------------- page numbering
def _page_field(paragraph):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = '1'
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    for el in (b, instr, sep, t, e):
        run._r.append(el)
    set_run_font(run, 12)


def configure_section(section, fmt=None, start=None, numbered=True):
    """Set page-number format/start and a centered PAGE field in the footer."""
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    if fmt is not None:
        pg = OxmlElement('w:pgNumType')
        pg.set(qn('w:fmt'), fmt)
        if start is not None:
            pg.set(qn('w:start'), str(start))
        sectPr.append(pg)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if numbered:
        _page_field(p)


# ---------------------------------------------------------------- TOC field
def insert_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = 'Nhấn Ctrl+A rồi F9 để cập nhật Mục lục.'
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar); run._r.append(instr)
    run._r.append(fldChar2); run._r.append(t); run._r.append(fldChar3)


def insert_tof():
    """Insert an auto Table of Figures (collects 'Hình' captions)."""
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\h \\z \\c "Hình"'
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = 'Nhấn Ctrl+A rồi F9 để cập nhật Danh mục hình.'
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    for el in (b, instr, sep, t, e):
        run._r.append(el)


# ---------------------------------------------------------------- table builder
def is_sep_row(cells):
    return all(re.match(r'^:?-{1,}:?$', c.strip()) for c in cells if c.strip() != '') and any(c.strip() for c in cells)


def split_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def add_table(rows):
    # rows: list of list-of-cells; first row header (sep row already removed)
    ncols = max(len(r) for r in rows)
    header = rows[0]
    body = rows[1:]
    tbl = doc.add_table(rows=1, cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hcells = tbl.rows[0].cells
    for j in range(ncols):
        shade_cell(hcells[j], TBL_HEADER_FILL)
        cp = hcells[j].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(2)
        txt = header[j] if j < len(header) else ''
        add_inline(cp, txt, 12, base_bold=True, base_color=BLACK)
    for ri, r in enumerate(body):
        cells = tbl.add_row().cells
        for j in range(ncols):
            cp = cells[j].paragraphs[0]
            cp.paragraph_format.space_after = Pt(2)
            cp.paragraph_format.line_spacing = 1.0
            txt = r[j] if j < len(r) else ''
            add_inline(cp, txt, 12)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------- figures
_mermaid_counter = [0]


def add_figure_caption(desc):
    """Centered italic 'Hình {SEQ}: desc' caption (collected by Table of Figures)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r1 = p.add_run('Hình ')
    set_run_font(r1, 12, bold=True, italic=True)
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = ' SEQ Hình \\* ARABIC '
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = '1'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    r2 = p.add_run()
    for el in (b, instr, sep, t, end):
        r2._r.append(el)
    set_run_font(r2, 12, bold=True, italic=True)
    r3 = p.add_run(': ' + desc)
    set_run_font(r3, 12, italic=True)


def _place_image(path):
    sz = png_size(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    if sz:
        ratio = sz[0] / sz[1]
        w_cm = MAX_W_CM
        h_cm = w_cm / ratio
        if h_cm > MAX_H_CM:
            h_cm = MAX_H_CM
            w_cm = h_cm * ratio
        run.add_picture(path, width=Cm(w_cm), height=Cm(h_cm))
    else:
        run.add_picture(path, width=Cm(MAX_W_CM))


def add_screenshot(path, caption):
    if not os.path.exists(path):
        # leave a visible placeholder so missing files are obvious
        para('[Thiếu ảnh: %s]' % os.path.basename(path),
             WD_ALIGN_PARAGRAPH.CENTER, 11, italic=True,
             color=RGBColor(0x99, 0x99, 0x99))
        return
    _place_image(path)
    add_figure_caption(caption)


def add_diagram_image():
    _mermaid_counter[0] += 1
    idx = _mermaid_counter[0]
    path = os.path.join(DIAG, 'mermaid_%02d.png' % idx)
    if not os.path.exists(path):
        return False
    _place_image(path)
    title = DIAGRAM_TITLES[idx - 1] if idx - 1 < len(DIAGRAM_TITLES) else 'Sơ đồ minh họa'
    add_figure_caption(title)
    return True


def add_code_box(lines, lang):
    if lang == 'mermaid':
        if add_diagram_image():
            return
        # fallback to code box if image missing
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    shade_paragraph(p, 'F4F4F4')
    add_box_border(p)
    text = '\n'.join(lines)
    r = p.add_run(text)
    set_run_font(r, 9, mono=True)


# ---------------------------------------------------------------- main parse
def main():
    raw = open(SRC, encoding='utf-8').read()
    raw = raw.replace('<![CDATA[', '').replace(']]>', '')
    lines = raw.split('\n')

    build_title_page()

    # Section 1 = title page: no page number.
    configure_section(doc.sections[0], numbered=False)
    # Section 2 = front matter, lowercase roman starting at i (new page).
    fm_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(fm_section, fmt='lowerRoman', start=1, numbered=True)

    i = 0
    n = len(lines)
    first_h1 = True   # first H1 (LỜI CẢM ƠN) already on a fresh section page

    # right-align block tracking
    right_align = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # skip the artificial "PHẦN MỞ ĐẦU" top marker
        if stripped == '# PHẦN MỞ ĐẦU':
            i += 1
            continue

        # div align right
        if stripped.lower().startswith('<div') and 'right' in stripped.lower():
            right_align = True; i += 1; continue
        if stripped.lower().startswith('</div'):
            right_align = False; i += 1; continue

        # horizontal rule -> ignore (acts as separator)
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            i += 1; continue

        # headings
        m = re.match(r'^(#{1,4})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            if level == 1:
                # LỜI MỞ ĐẦU starts the body section (arabic page numbers from 1)
                if re.match(r'^LỜI MỞ ĐẦU', title, re.I):
                    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
                    configure_section(body_section, fmt='decimal', start=1, numbered=True)
                    heading(title, 1, page_break=False)
                elif first_h1:
                    heading(title, 1, page_break=False)
                else:
                    heading(title, 1, page_break=True)
                first_h1 = False
                # MỤC LỤC -> insert auto TOC and skip following manual table
                if re.match(r'^MỤC LỤC', title, re.I):
                    insert_toc()
                    i += 1
                    while i < n and not re.match(r'^#\s+', lines[i]):
                        i += 1
                    continue
                # DANH MỤC HÌNH -> insert auto Table of Figures, skip stale list
                if re.match(r'^DANH MỤC HÌNH', title, re.I):
                    insert_tof()
                    i += 1
                    while i < n and not re.match(r'^#\s+', lines[i]):
                        i += 1
                    continue
            else:
                heading(title, level)
            i += 1
            continue

        # image:  ![caption](path)
        im = re.match(r'^!\[(.*?)\]\((.*?)\)\s*$', stripped)
        if im:
            alt, src = im.group(1).strip(), im.group(2).strip()
            if not os.path.isabs(src):
                src = os.path.join(BASE, src.replace('/', os.sep))
            add_screenshot(src, alt or 'Giao diện hệ thống')
            i += 1
            continue

        # table block
        if stripped.startswith('|') and '|' in stripped[1:]:
            tbl_lines = []
            while i < n and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i]); i += 1
            rows = [split_row(l) for l in tbl_lines]
            # remove separator rows
            rows = [r for r in rows if not is_sep_row(r)]
            if rows:
                add_table(rows)
            continue

        # fenced code
        fm = re.match(r'^```(\w*)', stripped)
        if fm:
            lang = fm.group(1)
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1  # skip closing fence
            add_code_box(code, lang)
            continue

        # blank
        if stripped == '':
            i += 1; continue

        # caption (italic line starting with *Bảng / *Hình)
        if re.match(r'^\*(Bảng|Hình)\s', stripped) and stripped.endswith('*'):
            para(stripped.strip('*'), WD_ALIGN_PARAGRAPH.CENTER, 12, italic=True, space_after=8)
            i += 1; continue

        # bullet list
        bm = re.match(r'^([*\-+])\s+(.*)$', line)
        if bm:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, bm.group(2))
            i += 1; continue

        # numbered list
        nm = re.match(r'^(\d+)\.\s+(.*)$', line)
        if nm:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, nm.group(2))
            i += 1; continue

        # blockquote
        if stripped.startswith('>'):
            p = para(stripped.lstrip('> ').strip(), size=12, italic=True,
                     color=RGBColor(0x55, 0x55, 0x55))
            i += 1; continue

        # normal paragraph (collect following non-empty, non-special lines? keep per-line)
        p = doc.add_paragraph()
        if right_align:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_inline(p, stripped)
        i += 1

    out_path = sys.argv[1] if len(sys.argv) > 1 else OUT
    doc.save(out_path)
    print('SAVED %s | paragraphs=%d tables=%d' % (out_path, len(doc.paragraphs), len(doc.tables)))


if __name__ == '__main__':
    main()
